from io import BytesIO

from docx import Document
from docx.shared import Inches, Pt
from reportlab.lib.colors import HexColor  # type: ignore[import-untyped]
from reportlab.lib.pagesizes import A4  # type: ignore[import-untyped]
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet  # type: ignore[import-untyped]
from reportlab.lib.units import mm  # type: ignore[import-untyped]
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer  # type: ignore[import-untyped]

from app.retrieval.models import CandidateProfile


def _lines(profile: CandidateProfile) -> list[tuple[str, str]]:
    return [
        ("Profile", profile.summary or "No summary provided."),
        ("Location", ", ".join(value for value in [profile.location, profile.country] if value)),
        ("Experience", f"{profile.years_experience:g} years" if profile.years_experience else ""),
        ("Age", f"{profile.age} years" if profile.age is not None else ""),
        ("Skills", ", ".join(profile.skills)),
        ("Technologies", ", ".join(profile.technologies)),
        ("Domains", ", ".join(profile.domains)),
        ("Companies", ", ".join(profile.companies)),
        ("Projects", ", ".join(profile.projects)),
    ]


def render_docx(profile: CandidateProfile) -> bytes:
    document = Document()
    section = document.sections[0]
    section.top_margin = section.bottom_margin = Inches(0.7)
    section.left_margin = section.right_margin = Inches(0.8)
    title = document.add_heading(profile.full_name, level=0)
    if title.style is not None:
        title.style.font.name = "Arial"
    subtitle = document.add_paragraph(profile.current_title or "Candidate")
    if subtitle.style is not None:
        subtitle.style.font.name = "Arial"
        subtitle.style.font.size = Pt(12)
    for label, value in _lines(profile):
        if not value:
            continue
        document.add_heading(label, level=1)
        paragraph = document.add_paragraph(value)
        if paragraph.style is not None:
            paragraph.style.font.name = "Arial"
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def render_pdf(profile: CandidateProfile) -> bytes:
    output = BytesIO()
    document = SimpleDocTemplate(
        output,
        pagesize=A4,
        leftMargin=18 * mm,
        rightMargin=18 * mm,
        topMargin=16 * mm,
        bottomMargin=16 * mm,
    )
    styles = getSampleStyleSheet()
    title = ParagraphStyle("CandidateTitle", parent=styles["Title"], textColor=HexColor("#12233f"))
    heading = ParagraphStyle(
        "CandidateHeading", parent=styles["Heading2"], textColor=HexColor("#1b6b5d")
    )
    story = [
        Paragraph(profile.full_name, title),
        Paragraph(profile.current_title or "Candidate", styles["Normal"]),
        Spacer(1, 8),
    ]
    for label, value in _lines(profile):
        if value:
            story.extend(
                [Paragraph(label, heading), Paragraph(value, styles["BodyText"]), Spacer(1, 6)]
            )
    document.build(story)
    return output.getvalue()
